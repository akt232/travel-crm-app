import streamlit as st
import pandas as pd
import plotly.express as px
import gspread
import json
import os
import re

from openai import OpenAI
from oauth2client.service_account import ServiceAccountCredentials
from datetime import datetime
from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2.service_account import Credentials
from PyPDF2 import PdfReader
import io
import re
# =====================================================
# CONFIG
# =====================================================

CONFIG_FILE = "config.json"

DEFAULT_SHEET = ""
DEFAULT_TOUR_SHEET = ""
DEFAULT_GUIDE_SHEET = "https://docs.google.com/spreadsheets/d/1b7z00QcNuYjK54ikc2ctbxsF3Ok7snGKSx57LChIZpA/edit#gid=0"
DRIVE_FOLDER_ID = ""   # folder chứa file tour trên Google Drive
LOGO_URL = "https://travel.com.vn/Content/images/logo.png"

st.set_page_config(
    page_title="Vietravel Sales Hub",
    page_icon="🌍",
    layout="wide"
)


# =====================================================
# LOAD CONFIG
# =====================================================

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)

    return {
        "sheet_url": DEFAULT_SHEET,
        "tour_sheet_url": DEFAULT_TOUR_SHEET,
        "guide_sheet_url": DEFAULT_GUIDE_SHEET,
        "api_key": ""
    }


def save_config(data):
    with open(CONFIG_FILE, "w") as f:
        json.dump(data, f, indent=4)


config = load_config()


# =====================================================
# SESSION
# =====================================================

if "api_key" not in st.session_state:
    st.session_state.api_key = config.get("api_key", "")

if "sheet_url" not in st.session_state:
    st.session_state.sheet_url = config.get("sheet_url", "")

if "tour_sheet_url" not in st.session_state:
    st.session_state.tour_sheet_url = config.get("tour_sheet_url", "")

if "guide_sheet_url" not in st.session_state:
    st.session_state.guide_sheet_url = config.get("guide_sheet_url", DEFAULT_GUIDE_SHEET)
if "drive_folder" not in st.session_state:
    st.session_state.drive_folder = config.get("drive_folder", "")

if "selected_customer" not in st.session_state:
    st.session_state.selected_customer = None

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "customer_list" not in st.session_state:
    st.session_state.customer_list = [
        {"id": 1, "name": "Anh Hùng", "msg": "Anh muốn đi Nhật tháng 3 ngân sách 40000000", "time": "10:30"},
        {"id": 2, "name": "Chị Lan", "msg": "Tour Thái Lan bao nhiêu tiền em?", "time": "09:15"},
        {"id": 3, "name": "Khách Web", "msg": "Tư vấn giúp tour Đà Nẵng", "time": "08:00"},
    ]


# =====================================================
# CSS
# =====================================================
st.markdown("""
<style>

/* ===== NỀN ===== */
html, body, .stApp {
    background: linear-gradient(135deg, #e0f2fe, #f8fbff) !important;
    color: #0f172a !important;
}

/* ===== SIDEBAR ===== */
[data-testid="stSidebar"] {
    background: #dbeafe !important;
}

[data-testid="stSidebar"] * {
    color: #0f172a !important;
    font-weight: 500;
}

/* ===== TEXT GLOBAL ===== */
h1, h2, h3, h4, h5, h6, p, span, label, div {
    color: #0f172a !important;
}

/* ===== INPUT ===== */
.stTextInput input,
.stTextArea textarea {
    background: white !important;
    color: #0f172a !important;
    border: 1px solid #cbd5e1 !important;
}

/* ===== BUTTON ===== */
.stButton>button {
    background: #2563eb !important;
    color: white !important;
    border-radius: 8px;
    border: none;
    height: 40px;
    font-weight: 600;
}

.stButton>button:hover {
    background: #1d4ed8 !important;
}

/* ===== CHAT BOX ===== */
.chat-box {
    background: white !important;
    border: 1px solid #cbd5e1;
    border-radius: 12px;
    height: 60vh;
    display: flex;
    flex-direction: column;
    box-shadow: 0 4px 15px rgba(0,0,0,0.05);
}

.chat-area {
    flex-grow: 1;
    overflow-y: auto;
    padding: 15px;
}

/* ===== MESSAGE ===== */
.msg {
    background: #e0f2fe;
    padding: 10px;
    border-radius: 8px;
    margin-bottom: 10px;
    color: #0f172a;
}

/* ===== METRIC BOX ===== */
[data-testid="metric-container"] {
    background: white;
    border: 1px solid #e2e8f0;
    padding: 10px;
    border-radius: 10px;
}

/* ===== PLACEHOLDER TEXT ===== */
::placeholder {
    color: #64748b !important;
}

/* ===== RADIO / SELECT ===== */
.stSelectbox div,
.stRadio label {
    color: #0f172a !important;
}

/* ===== EXPANDER ===== */
.streamlit-expanderHeader {
    color: #0f172a !important;
    font-weight: 600;
}

/* ===== REMOVE DARK OVERLAY ===== */
[data-testid="stHeader"] {
    background: transparent !important;
}

/* =========================
   SELECTBOX MAIN
========================= */

.stSelectbox div[data-baseweb="select"] > div {
    background: linear-gradient(135deg, #3b82f6, #2563eb) !important;
    color: white !important;
    border-radius: 10px !important;
    border: 1px solid #2563eb !important;
}


/* Text trong select */
.stSelectbox span {
    color: white !important;
    font-weight: 500;
}


/* Dropdown menu */
div[data-baseweb="popover"] {
    background: #1e40af !important;
    border-radius: 10px !important;
}


/* Item trong dropdown */
div[role="option"] {
    background: #1e40af !important;
    color: white !important;
}


/* Hover item */
div[role="option"]:hover {
    background: #2563eb !important;
    color: white !important;
}


/* Remove viền đỏ focus */
.stSelectbox div[data-baseweb="select"]:focus-within {
    box-shadow: 0 0 0 2px #60a5fa !important;
    border-color: #60a5fa !important;
}


/* Icon dropdown */
.stSelectbox svg {
    fill: white !important;
}

</style>
""", unsafe_allow_html=True)
st.markdown("""
<style>

/* =========================
   SELECTBOX MAIN BOX
========================= */

.stSelectbox div[data-baseweb="select"] > div {
    background: linear-gradient(135deg, #3b82f6, #2563eb) !important;
    color: white !important;
    border-radius: 10px !important;
    border: 1px solid #2563eb !important;
}


/* Text selected */
.stSelectbox div[data-baseweb="select"] span {
    color: white !important;
}


/* Icon dropdown */
.stSelectbox svg {
    fill: white !important;
}


/* =========================
   DROPDOWN MENU FIX ĐEN
========================= */

/* Menu container */
div[data-baseweb="popover"],
div[data-baseweb="menu"] {
    background: #1e40af !important;
    color: white !important;
}


/* Option item */
div[role="option"] {
    background: #1e40af !important;
    color: white !important;
}


/* Hover option */
div[role="option"]:hover {
    background: #2563eb !important;
    color: white !important;
}


/* Selected option */
div[aria-selected="true"] {
    background: #3b82f6 !important;
    color: white !important;
}


/* Remove nền đen sâu bên trong */
ul, li {
    background: transparent !important;
}


/* Focus border */
.stSelectbox div[data-baseweb="select"]:focus-within {
    border-color: #60a5fa !important;
    box-shadow: 0 0 0 2px #60a5fa !important;
}

</style>
""", unsafe_allow_html=True)
# =====================================================
# CHATGPT FUNCTION
# =====================================================

def ask_chatgpt(prompt):
    if not st.session_state.api_key:
        return "Chưa nhập OpenAI API Key"

    try:
        client = OpenAI(api_key=st.session_state.api_key)

        response = client.chat.completions.create(
            model="gpt-4o-mini",  # Đã sửa từ gpt-4.1-mini thành gpt-4o-mini
            messages=[
                {"role": "system", "content": "Bạn là chuyên gia du lịch."},
                {"role": "user", "content": prompt}
            ]
        )

        return response.choices[0].message.content

    except Exception as e:
        return str(e)
# =====================================================
# GOOGLE SHEET
# =====================================================

def connect_sheet(url):

    scope = [
        "https://spreadsheets.google.com/feeds",
        "https://www.googleapis.com/auth/drive"
    ]

    creds_dict = st.secrets["gcp_service_account"]

    creds = ServiceAccountCredentials.from_json_keyfile_dict(
        creds_dict,
        scope
    )

    client = gspread.authorize(creds)

    sheet = client.open_by_url(url).sheet1

    return sheet


def load_sheet():
    try:
        sheet = connect_sheet(st.session_state.sheet_url)
        data = sheet.get_all_records()
        return pd.DataFrame(data)
    except:
        return pd.DataFrame()


def load_tour_sheet():
    try:
        sheet = connect_sheet(st.session_state.tour_sheet_url)
        data = sheet.get_all_records()
        return pd.DataFrame(data)
    except:
        return pd.DataFrame()


def load_guide_sheet(worksheet_name=None):
    try:
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds_dict = st.secrets["gcp_service_account"]
        creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
        client = gspread.authorize(creds)
        spreadsheet = client.open_by_url(st.session_state.guide_sheet_url)
        
        # Nếu có tên worksheet thì mở, không thì mở sheet đầu tiên
        sheet = spreadsheet.worksheet(worksheet_name) if worksheet_name else spreadsheet.sheet1
        
        data = sheet.get_all_records()
        return pd.DataFrame(data)
    except Exception as e:
        st.error(f"Lỗi: {e}")
        return pd.DataFrame()
def get_guide_worksheets():
    try:
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds_dict = st.secrets["gcp_service_account"]
        creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
        client = gspread.authorize(creds)
        spreadsheet = client.open_by_url(st.session_state.guide_sheet_url)
        return [sh.title for sh in spreadsheet.worksheets()]
    except:
        return []
def save_to_sheet(row):
    try:
        sheet = connect_sheet(st.session_state.sheet_url)
        sheet.append_row(row)
        return True
    except Exception as e:
        st.error(e)
        return False


def delete_row(row_number):
    try:
        sheet = connect_sheet(st.session_state.sheet_url)
        sheet.delete_rows(row_number)
        return True
    except:
        return False

import io
import streamlit as st

from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2.service_account import Credentials

from PyPDF2 import PdfReader
from docx import Document


# =============================
# CONFIG
# =============================




# =============================
# CONNECT GOOGLE DRIVE
# =============================

def connect_drive():

    scope = ["https://www.googleapis.com/auth/drive.readonly"]

    creds_dict = st.secrets["gcp_service_account"]

    creds = Credentials.from_service_account_info(
        creds_dict,
        scopes=scope
    )

    service = build("drive", "v3", credentials=creds)

    return service


# =============================
# READ PDF
# =============================

def read_pdf_from_bytes(file_bytes):

    pdf = PdfReader(file_bytes)
    text = ""

    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text


# =============================
# READ DOCX
# =============================

def read_docx_from_bytes(file_bytes):

    doc = Document(file_bytes)

    text = []
    for p in doc.paragraphs:
        text.append(p.text)

    return "\n".join(text)


# =============================
# LOAD ALL TOUR DATA FROM DRIVE
# =============================

def extract_drive_id(link):
    import re
    match = re.search(r'/folders/([a-zA-Z0-9_-]+)', link)
    return match.group(1) if match else link


def load_drive_tour_data():

    # ===== LẤY FOLDER ID TỪ SESSION =====
    drive_link = st.session_state.get("drive_folder", "")

    if not drive_link:
        st.warning("⚠️ Chưa cấu hình Google Drive Folder trong Settings.")
        return ""

    folder_id = extract_drive_id(drive_link)

    try:

        service = connect_drive()

        results = service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields="files(id, name, mimeType)",
            pageSize=100
        ).execute()

        files = results.get("files", [])

        if len(files) == 0:
            st.warning("⚠️ Folder có nhưng không có file hoặc chưa share quyền.")
            return ""

        all_text = ""

        for file in files:

            file_id = file["id"]
            file_name = file["name"].lower()

            try:

                request = service.files().get_media(fileId=file_id)

                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)

                done = False
                while not done:
                    status, done = downloader.next_chunk()

                fh.seek(0)

                if file_name.endswith(".pdf"):
                    text = read_pdf_from_bytes(fh)

                elif file_name.endswith(".docx"):
                    text = read_docx_from_bytes(fh)

                elif file_name.endswith(".txt"):
                    text = fh.read().decode("utf-8")

                else:
                    continue

                all_text += "\n" + text

            except Exception as e:
                st.error(f"Lỗi đọc file {file_name}: {e}")

        return all_text

    except Exception as e:
        st.error(f"Lỗi kết nối Drive: {e}")
        return ""

# =============================
# AI SEARCH TOUR FROM DRIVE DATA
# =============================
def search_relevant_text(data, query, window=8000):

    query = query.lower()

    idx = data.lower().find(query)

    if idx == -1:
        return data[:window]

    start = max(0, idx - window)
    end = idx + window

    return data[start:end]
def ai_search_tour_drive(query):

    data = load_drive_tour_data()

    if not data:
        return "❌ Không có dữ liệu Drive. Vui lòng kiểm tra Folder ID hoặc quyền chia sẻ."

    # =============================
    # TÌM ĐOẠN LIÊN QUAN NHẤT
    # =============================
    relevant = search_relevant_text(data, query, window=9000)

    # =============================
    # PROMPT CHUẨN PRO
    # =============================
    prompt = f"""
Bạn là chuyên gia sản phẩm Vietravel.

NHIỆM VỤ:
Trích xuất CHÍNH XÁC thông tin tour từ dữ liệu được cung cấp.

=============================
QUY TẮC BẮT BUỘC
=============================

1. CHỈ sử dụng dữ liệu có trong tài liệu
2. KHÔNG được tự thêm thông tin ngoài dữ liệu
3. Nếu không thấy thông tin thì ghi: Đang cập nhật
4. Phải hiển thị ĐẦY ĐỦ tất cả các ngày trong lịch trình
5. Nếu tour 7 ngày phải có Ngày 1 → Ngày 7
6. Không được bỏ sót ngày cuối
7. Không được tóm tắt quá ngắn
8. Ưu tiên dữ liệu gần từ khóa tìm kiếm: "{query}"
9. Viết văn phong tư vấn chuyên nghiệp gửi khách hàng
10. Nội dung phải dài và đầy đủ

=============================
DỮ LIỆU TOUR
=============================

{relevant}

=============================
KHÁCH HỎI
=============================

{query}

=============================
XUẤT KẾT QUẢ THEO FORMAT
=============================

📍 Tên tour:
📍 Mã tour:
📍 Thời gian:
📍 Giá:
📍 Ngày khởi hành:

📍 Điểm nổi bật:

📍 LỊCH TRÌNH CHI TIẾT:

Ngày 1:
Ngày 2:
Ngày 3:
Ngày 4:
(Nếu còn ngày phải liệt kê đầy đủ đến ngày cuối)

=============================

Viết nội dung rõ ràng dễ copy gửi Zalo cho khách.
Không được rút gọn.
"""

    result = ask_chatgpt(prompt)

    return result
# =====================================================
# TOUR SUGGEST
# =====================================================

STOP_WORDS = [
    "tư", "vấn", "giúp", "tour", "muốn", "đi", "em", "anh",
    "chị", "bao", "nhiêu", "tiền", "tháng", "ngân", "sách"
]


def clean_words(text):
    words = re.findall(r'\w+', text.lower())
    return [w for w in words if w not in STOP_WORDS and len(w) > 2]


def suggest_tour(message):

    df = load_tour_sheet()

    if df.empty:
        return pd.DataFrame()

    msg = message.lower()

    STOP_WORDS = [
        "anh", "chị", "em", "mình", "tôi",
        "muốn", "đi", "du", "lịch", "tour",
        "tháng", "ngày", "bao", "nhiêu",
        "tiền", "ngân", "sách", "khoảng",
        "tầm", "giúp", "với", "ạ", "ơi"
    ]

    words = re.findall(r'\w+', msg)

    keywords = [
        w for w in words
        if w not in STOP_WORDS and len(w) > 2 and not w.isdigit()
    ]

    if not keywords:
        return pd.DataFrame()

    results = []

    for _, row in df.iterrows():

        # CHỈ LẤY TÊN TOUR ĐỂ SO SÁNH
        tour_name = str(row.get("Tour (Tên tour)", "")).lower()

        matched = False

        for kw in keywords:
            if kw in tour_name:
                matched = True
                break

        if matched:
            results.append(row)

    if not results:
        return pd.DataFrame()

    return pd.DataFrame(results)


# =====================================================
# DASHBOARD
# =====================================================

def render_dashboard():

    st.title("📊 Dashboard")

    df = load_sheet()

    if df.empty:
        st.warning("Chưa có dữ liệu")
        return

    # ===== CLEAN DATA =====
    if "Giá" in df.columns:
        df["Giá"] = (
            df["Giá"]
            .astype(str)
            .str.replace(",", "", regex=False)
            .str.replace("đ", "", regex=False)
        )
        df["Giá"] = pd.to_numeric(df["Giá"], errors="coerce").fillna(0)

    if "Ngày" in df.columns:
        df["Ngày"] = pd.to_datetime(df["Ngày"], errors="coerce")

    # ===== TODAY (GIỜ VIỆT NAM) =====
    from datetime import datetime, timedelta

    vietnam_now = datetime.utcnow() + timedelta(hours=7)
    today = vietnam_now.date()

    if "Ngày" in df.columns:
        today_df = df[df["Ngày"].dt.date == today]
    else:
        today_df = pd.DataFrame()

    # ===== METRICS =====
    today_customers = len(today_df)
    today_revenue = today_df["Giá"].sum()

    total_customers = len(df)
    total_revenue = df["Giá"].sum()

    col1, col2, col3, col4 = st.columns(4)

    col1.metric("Khách hôm nay", today_customers)
    col2.metric("Doanh thu hôm nay", f"{today_revenue:,.0f} đ")
    col3.metric("Tổng khách", total_customers)
    col4.metric("Tổng doanh thu", f"{total_revenue:,.0f} đ")

    st.divider()

    # ===== DOANH THU THEO TOUR =====
    if "Tour" in df.columns:

        route_df = df.groupby("Tour").agg({
            "Tên": "count",
            "Giá": "sum"
        }).reset_index()

        fig1 = px.bar(
            route_df,
            x="Tour",
            y="Giá",
            color="Tour",
            title="Doanh thu theo Tour"
        )

        st.plotly_chart(fig1, use_container_width=True)

    # ===== DOANH THU THEO NGÀY =====
    if "Ngày" in df.columns:

        daily = df.groupby(df["Ngày"].dt.date)["Giá"].sum().reset_index()

        fig2 = px.line(
            daily,
            x="Ngày",
            y="Giá",
            markers=True,
            title="Doanh thu theo ngày"
        )

        st.plotly_chart(fig2, use_container_width=True)

# =====================================================
# SALES CENTER
# =====================================================

def render_sales_center():

    col_left, col_mid, col_right = st.columns([1, 2, 1])

    # ================= LEFT =================
    with col_left:

        st.subheader("Khách hàng")

        for cust in st.session_state.customer_list:
            if st.button(f"{cust['name']} - {cust['time']}", key=cust["id"]):
                st.session_state.selected_customer = cust

    # ================= MID =================
    with col_mid:

        cust = st.session_state.selected_customer

        if cust:

            st.subheader(f"Chat với {cust['name']}")

            st.markdown(f"""
            <div class="chat-box">
                <div class="chat-area">
                    <div class="msg">{cust["msg"]}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # ===== TOUR SUGGEST =====
            st.subheader("🎯 Tour phù hợp")

            suggest_df = suggest_tour(cust["msg"])

            if suggest_df.empty:
                st.info("Không tìm thấy tour")
            else:
                st.dataframe(suggest_df)

            # ===== AI REPLY =====
            st.subheader("🤖 AI gợi ý trả lời (theo dữ liệu công ty)")

            if st.button("Gợi ý trả lời khách"):
                prompt = f"Khách nói: {cust['msg']}. Hãy trả lời tư vấn tour chuyên nghiệp."
                reply = ask_company_ai(prompt)
                st.success(reply)

            # ===== AI OBJECTION =====
            st.subheader("🧠 Xử lý từ chối")

            if st.button("Gợi ý xử lý từ chối"):
                prompt = f"""
Khách nói: {cust['msg']}

Đưa ra 3 cách xử lý chuyên nghiệp để thuyết phục khách.
"""
                reply = ask_chatgpt(prompt)
                st.info(reply)

            # ===== STATUS =====
            status = st.selectbox(
                "Trạng thái",
                ["Đang theo dõi", "Đã chốt đơn", "Không chốt"]
            )

            if status == "Đã chốt đơn":

                with st.form("deal"):

                    name = st.text_input("Tên", cust["name"])
                    tour = st.text_input("Tour")
                    price = st.text_input("Giá")
                    note = st.text_area("Note")
                    sale = st.text_input("Sale")

                    channel = st.selectbox(
                        "Kênh",
                        ["Online", "Facebook", "Zalo", "Chi nhánh"]
                    )

                    ok = st.form_submit_button("Xác nhận")

                    if ok:

                        saved = save_to_sheet([
                            datetime.now().strftime("%Y-%m-%d"),
                            name,
                            tour,
                            price,
                            note,
                            channel,
                            sale
                        ])

                        if saved:
                            st.success("✅ Đã lưu Google Sheet")

    # ================= RIGHT =================
    with col_right:

        # ===== AI TRA CỨU NỘI BỘ =====
        st.subheader("⚡ AI Tra cứu nội bộ")

        user_q = st.text_input("Hỏi dữ liệu công ty")

        if st.button("Tra cứu"):

            res = ask_company_ai(user_q)

            st.session_state.chat_history.append(("Bạn", user_q))
            st.session_state.chat_history.append(("AI", res))

        # ===== AI SO SÁNH TOUR =====
        st.subheader("📊 So sánh 2 tour")

        tour1 = st.text_input("Tour 1")
        tour2 = st.text_input("Tour 2")

        if st.button("So sánh tour"):

            prompt = f"So sánh 2 tour {tour1} và {tour2} của công ty Vietravel."

            res = ask_company_ai(prompt)

            st.session_state.chat_history.append(("Bạn", f"So sánh: {tour1} vs {tour2}"))
            st.session_state.chat_history.append(("AI", res))

        # =============================
        # AI TRA TOUR DRIVE (NEW)
        # =============================

        st.divider()
        st.subheader("📂 AI Tra cứu Tour (Drive)")

        drive_query = st.text_input(
            "Nhập tên tour cần tìm",
            placeholder="Ví dụ: Nhật Bản, Hàn Quốc, Úc..."
        )

        if st.button("🔍 Tìm Tour Drive"):

            if not drive_query:
                st.warning("Nhập tên tour")
            else:

                with st.spinner("AI đang đọc dữ liệu Drive..."):

                    result = ai_search_tour_drive(drive_query)

                st.success("✅ Đã tìm thấy thông tin")

                st.text_area(
                    "Thông tin gửi khách",
                    result,
                    height=300
                )

                # COPY BOX
                st.code(result, language="text")

        # ===== CHAT HISTORY =====
        st.subheader("💬 Lịch sử AI")

        for role, msg in st.session_state.chat_history:
            st.write(f"**{role}:** {msg}")

# =====================================================
# CUSTOMERS & ORDERS
# =====================================================

def render_customer_orders():

    st.title("Customers & Orders")

    st.subheader("Danh sách khách")
    st.dataframe(pd.DataFrame(st.session_state.customer_list))

    st.divider()

    df = load_sheet()

    st.subheader("Đơn đã chốt")

    if df.empty:
        st.info("Chưa có dữ liệu")
        return

    for idx, row in df.iterrows():

        col1, col2, col3, col4, col5, col6 = st.columns([2,2,2,2,2,1])

        with col1:
            st.write(row.get('Ngày',''))

        with col2:
            st.write(row.get('Tên',''))

        with col3:
            st.write(row.get('Tour',''))

        with col4:
            st.write(row.get('Giá',''))

        with col5:
            st.write(row.get('Kênh',''))

        with col6:

            if st.button("❌", key=f"del_{idx}"):

                ok = delete_row(idx + 2)

                if ok:
                    st.success("Đã xóa")
                    st.rerun()


# =====================================================
# GUIDE CENTER
# =====================================================

# =====================================================
# GUIDE CENTER
# =====================================================

# =====================================================
# GUIDE CENTER
# =====================================================
def render_guide_center():

    st.title("📘 Cẩm nang")

    # =========================
    # DATA MENU FULL
    # =========================

    guide_data = {

        "CÁC YÊU CẦU": [
            ("Check đối thủ T5 hàng tuần", "https://drive.google.com/drive/folders/1CjjaWQ6AXM-gh70s6MXfQHLYmhdsQ6gY"),
            ("Ticket E-com - Quầy", "https://gemini.google.com/share/75ec867c7a43"),
            ("Link lưu trữ hồ sơ", "https://docs.google.com/spreadsheets/d/1hkvZYjtHQjWATXlICl2Dhw7UHFTaViQx/edit?gid=359318821#gid=359318821"),
        ],

        "CÁC BÁO CÁO": [
            ("BC nhu cầu khách quan tâm tour hàng ngày", "https://docs.google.com/spreadsheets/d/1xN38cSycrwYpKd0ho3O7zhVtP6zEIjdXW03kxBSseuc/edit"),
            ("Báo cáo chi phí chi nhánh ghép", "https://docs.google.com/spreadsheets/d/1rcQgEkYsD46B4Wk1ZzcmgHVZfVsyqcSG/edit?gid=1378604091#gid=1378604091"),
            ("Báo cáo khách bị từ chối visa", "https://docs.google.com/spreadsheets/d/102pizHsZK-dXdqaz8LG_48dh5Phiqe8A/edit?gid=1735879678#gid=1735879678"),
            ("Báo cáo phát sinh VU 2025", "https://docs.google.com/spreadsheets/d/1vZTugGe1QXLwQhy9bxUGpkKjmD3YInLm8D8F_8Lo6jc/edit"),
            ("Đánh giá kênh PR online", "https://drive.google.com/drive/folders/1j0wakn6HFirnv4by5-77Nopjkh6pkaj_"),
            ("Báo cáo tình hình khách theo sự vụ 2025", "https://docs.google.com/spreadsheets/d/1X8f_VB5zsA65YCytpVPZ4ymyOidAVFQM_uJ_n2LJfuI/edit?gid=0#gid=0"),
            ("Báo cáo khách phản ánh HDV", "https://docs.google.com/spreadsheets/d/1ny4MgIOY9oS5VZFAmEDfFW8bp0KCbdcN/edit?gid=1378604091#gid=1378604091"),
            ("Báo cáo khách hủy tour theo thị trường", "https://docs.google.com/spreadsheets/d/1kuy-bE4hbGgfBDArI55MYDcG4kQGquNf/edit?gid=247600996#gid=247600996"),
            ("Khảo sát nhu cầu khách hàng", "https://docs.google.com/spreadsheets/d/1Zcv7EXGzT-urIHpMOTvdFdcUgx16vkS1/edit?gid=283996237#gid=283996237"),
            ("Kế hoạch Telesale", "http://drive.google.com/drive/folders/1aK1fv01nFGoD1MkTRtv93uiDYJ2PgeYl"),
        ],

        "KH TRUYỀN THÔNG 2025": [
            ("Lịch đăng bài Fanpage Vietravel hàng tuần", "https://docs.google.com/spreadsheets/d/15oTJTvTa95SREy66MR6SnsLhg8QtWFhT-ozR7DKIXms/edit?gid=1127315367#gid=1127315367"),
        ],

        "THÔNG TIN CHUNG": [
            ("Danh sách đăng ký hoàn tiền cho khách", "https://docs.google.com/spreadsheets/d/1p3ugCM9ZkQqZad9tNAyfXkSvMWKSN1AI/edit"),
            ("Thông tin họp đoàn mẫu", "https://drive.google.com/drive/folders/1aodfsPrA6ey0yeCe_byLAxINtzYOBP1t"),
            ("Tạo QR code", "https://docs.google.com/spreadsheets/u/0/d/1u0U_cEHTM1doUihaRXsH8bVtyIaDNgzF/edit"),
            ("Tổng hợp các chương trình ưu đãi tài chính", "https://docs.google.com/spreadsheets/d/12-X-UHL2SeyHgWbGeRxZHrIrsdNPDKom/edit?gid=11240248#gid=11240248"),
            ("Video Sản Phẩm của BSP", "https://drive.google.com/drive/folders/1ssuW6KhIgfDhlNsf8-2Xjxxp8sD8gfQN"),
            ("Chi tiết mức chi Hoa Hồng", "https://docs.google.com/spreadsheets/d/1vxIyiM04-7HM7l5Tg_jtWQnhlrpYw74N/edit?gid=1266470926#gid=1266470926"),
            ("Khuyến mãi Xuân 2026", "https://docs.google.com/spreadsheets/d/1hgMkuFrWoktsKKn-dma23pq0tIS6v2BJyL5yTO0IWu8/edit?gid=0#gid=0"),
            ("Tổng hợp quy trình Trung tâm FIT", "https://docs.google.com/spreadsheets/d/1zbF49Pa1Eq7dX1X7iop-lfvi32Wk_rMeLSvgc2DWEn8/edit?gid=0#gid=0"),
            ("TTX take note nhắc nhở", "https://docs.google.com/spreadsheets/d/1UrlqAxcXZBbfLTBZoFcB4pFNXhPRv3mrpjerS2J1IZg/edit?gid=0#gid=0"),
        ],

        "CẨM NANG TƯ VẤN DV SP BÁN": [
            ("Thông tin tour Châu Âu", "https://docs.google.com/spreadsheets/d/1ELv_B6EuRLO9But48qlYFQ_RPuftsHTw/edit?gid=1277256155#gid=1277256155"),
            ("Thông tin tour Châu Mỹ", "https://docs.google.com/spreadsheets/d/1AcEYg9sVKYPbz__MCrLdflHUr1ExG3SmAP1ioo_RnHU/edit?gid=0#gid=0"),
            ("Thông tin tour Châu Úc", "https://docs.google.com/spreadsheets/d/1-U4XjUw44buIgXye4MGh4Oqdsars64hZ/edit?gid=1277256155#gid=1277256155"),
            ("Thông tin tour Nhật Bản", "https://docs.google.com/spreadsheets/d/17qLBrdl6Wrz5MfDeCzR31QAEjciuQhgP/edit?gid=566600147#gid=566600147"),
            ("Thông tin tour Hàn Quốc", "https://docs.google.com/spreadsheets/d/1pWJ4igSPzsG7kSGpUiadTODCFnP1jsi5XRXD17ZkSc0/edit?gid=1413342922#gid=1413342922"),
            ("Thông tin tour Thái Lan", "https://docs.google.com/spreadsheets/d/1rDKIPnDIsgaNxWmCI3rGtJDCGiysRybzh7xrIO3t6js/edit?gid=0#gid=0"),
            ("Thông tin tour Tiếng Hoa", "https://drive.google.com/drive/folders/1x6UAxej421ujrStmAi3t40nwKKo9BGaE?usp=drive_link"),
            ("Thông tin tour Đông Nam Á (trừ Thái Lan)", "https://docs.google.com/spreadsheets/d/1lPTTo-Scd2BNPImpSvmtUrO5GQPH5ijq/edit?gid=124368584#gid=124368584"),
            ("Thông tin tour Miền Bắc", "https://docs.google.com/spreadsheets/d/1PlLTLEgZCg6QKYLFY5-vs_65ZtgzP56ts03Orw_Cx7Y/edit?gid=0#gid=0"),
            ("Thông tin tour Miền Trung", "https://docs.google.com/spreadsheets/d/1GaUlMYm3f_col5wQ_kaV2n9X0zqXiyKsK2O74Rn84Bo/edit?gid=1461750039#gid=1461750039"),
            ("Thông tin tour Miền Nam", "https://docs.google.com/spreadsheets/d/1L446JdJtmcaFZRS0kpMvlAq7P4aJNWZwFDUUFJJpEUE/edit?gid=0#gid=0"),
            ("Đào tạo nội bộ FIT - Google Drive", "https://drive.google.com/drive/folders/1jzs5xZfKJZGX_6pibuz2JiaL9eTeiB7g"),
        ],
    }

    # =========================
    # MENU LỚN
    # =========================

    st.subheader("Chọn mục")

    cols = st.columns(3)

    for i, category in enumerate(guide_data.keys()):
        with cols[i % 3]:
            if st.button(category, use_container_width=True):
                st.session_state["guide_category"] = category

    if "guide_category" not in st.session_state:
        st.session_state["guide_category"] = list(guide_data.keys())[0]

    selected_category = st.session_state["guide_category"]

    st.divider()

    st.subheader(selected_category)

    # =========================
    # MENU NHỎ
    # =========================

    items = guide_data[selected_category]

    for name, link in items:
        st.link_button(
            f"📄 {name}",
            link,
            use_container_width=True
        )
# =====================================================
# VISA AI
# =====================================================

def read_docx(file_path):
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except:
        return ""


visa_rule_1 = read_docx("THÔNG BÁO NHẬN QT NN.docx")
visa_rule_2 = read_docx("CÁC LƯU Ý VISA NHẬP CẢNH VIỆT NAM CHO NGƯỜI NƯỚC NGOÀI.docx")

visa_knowledge = visa_rule_1 + "\n" + visa_rule_2


def visa_tab():

    st.title("🛂 Visa Information")

    nationality = st.text_input("Quốc tịch")
    destination = st.text_input("Điểm đến")

    if st.button("Kiểm tra Visa"):

        prompt = f"""
Dữ liệu:
{visa_knowledge}

Khách quốc tịch {nationality} đi {destination}.

Tư vấn visa chi tiết.
"""

        result = ask_chatgpt(prompt)
        st.write(result)

# =====================================================
# COMPANY AI KNOWLEDGE BASE
# =====================================================

def load_company_knowledge():

    text = ""

    # Visa docs
    text += visa_knowledge + "\n"

    # Tour sheet
    try:
        df = load_tour_sheet()
        if not df.empty:
            text += df.to_string()
    except:
        pass

    return text


def ask_company_ai(question):

    knowledge = load_company_knowledge()

    prompt = f"""
Bạn là chuyên gia sản phẩm Vietravel.

Dữ liệu nội bộ công ty:
{knowledge}

Câu hỏi:
{question}

Trả lời chính xác theo dữ liệu công ty.
"""

    return ask_chatgpt(prompt)
# =====================================================
# SETTINGS
# =====================================================

def render_settings():

    st.title("Settings")

    key = st.text_input(
        "OpenAI API Key",
        value=st.session_state.api_key,
        type="password"
    )

    if st.button("Save API"):

        st.session_state.api_key = key

        save_config({
            "sheet_url": st.session_state.sheet_url,
            "tour_sheet_url": st.session_state.tour_sheet_url,
            "guide_sheet_url": st.session_state.guide_sheet_url,
            "drive_folder": st.session_state.get("drive_folder", ""),
            "api_key": key
        })

        st.success("Saved permanently ✅")

    st.divider()

    sheet_link = st.text_input(
        "Link Sheet Orders",
        value=st.session_state.sheet_url
    )

    tour_link = st.text_input(
        "Link Sheet Tour",
        value=st.session_state.tour_sheet_url
    )

    guide_link = st.text_input(
        "Link Sheet Guide",
        value=st.session_state.guide_sheet_url
    )

    # ===============================
    # NEW — DRIVE TOUR FOLDER
    # ===============================

    if "drive_folder" not in st.session_state:
        st.session_state.drive_folder = config.get("drive_folder", "")

    drive_link = st.text_input(
        "📂 Link Google Drive Folder (Tour Files)",
        value=st.session_state.drive_folder,
        placeholder="Dán link folder Google Drive chứa file tour..."
    )

    if st.button("Lưu cấu hình"):

        st.session_state.sheet_url = sheet_link
        st.session_state.tour_sheet_url = tour_link
        st.session_state.guide_sheet_url = guide_link
        st.session_state.drive_folder = drive_link

        save_config({
            "sheet_url": sheet_link,
            "tour_sheet_url": tour_link,
            "guide_sheet_url": guide_link,
            "drive_folder": drive_link,
            "api_key": st.session_state.api_key
        })

        st.success("Đã lưu vĩnh viễn")
# =====================================================
# SIDEBAR
# =====================================================

st.sidebar.image(LOGO_URL, width=150)

menu = st.sidebar.radio(
    "MENU",
    ["Dashboard", "Sales Center", "Customers & Orders", "Guide Center", "Visa Info", "Settings"]
)


# =====================================================
# ROUTER
# =====================================================

if menu == "Dashboard":
    render_dashboard()

elif menu == "Sales Center":
    render_sales_center()

elif menu == "Customers & Orders":
    render_customer_orders()

elif menu == "Guide Center":
    render_guide_center()

elif menu == "Visa Info":
    visa_tab()

elif menu == "Settings":
    render_settings()

























